"""Build a teacher-facing project report and a separate speaking script."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_final_oral_report_docx as base


ROOT = Path(__file__).resolve().parents[1]
REPORT_OUTPUT = ROOT / "docs" / "陈中远_孪生光储备池阶段工作成果报告_20260731.docx"
SCRIPT_OUTPUT = ROOT / "docs" / "陈中远_阶段工作成果报告讲解稿_20260731.docx"


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _set_running_furniture(section, left_text: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(left_text)
    base._set_run_font(run, size=8.5, color=base.MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    _clear_paragraph(paragraph)
    base._add_page_number(paragraph)


def _cover(
    doc: Document,
    *,
    kicker: str,
    title: str,
    subtitle: str,
    rows: list[list[str]],
    takeaway: str,
) -> None:
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(11)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(kicker)
    base._set_run_font(run, size=12, bold=True, color=base.BLUE)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title)
    base._set_run_font(run, size=27, bold=True, color=base.DARK_BLUE)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(32)
    run = paragraph.add_run(subtitle)
    base._set_run_font(run, size=15, color=base.MUTED)

    base._add_table(
        doc,
        ["项目", "内容"],
        rows,
        [2100, 7260],
        first_col_left=False,
    )
    base._add_callout(
        doc,
        "核心结论：",
        takeaway,
        fill=base.LIGHT_BLUE,
        accent=base.DARK_BLUE,
    )
    doc.add_page_break()


def _set_image_alt(last_paragraph, description: str) -> None:
    drawings = last_paragraph._p.xpath(".//w:drawing")
    if not drawings:
        return
    doc_pr = drawings[-1].xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", description)


def _add_report_figure(
    doc: Document,
    path: Path,
    caption: str,
    description: str,
    width: float,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    _set_image_alt(paragraph, description)

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(2)
    caption_p.paragraph_format.space_after = Pt(9)
    caption_run = caption_p.add_run(caption)
    base._set_run_font(caption_run, size=9, italic=True, color=base.MUTED)


def _apply_compact_guide_styles(doc: Document) -> None:
    """Resolve the compact_reference_guide preset for the speaking notes."""
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, base.BLUE, 18, 10),
        "Heading 2": (13, base.BLUE, 14, 7),
        "Heading 3": (12, base.DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_report() -> Path:
    doc = Document()
    base._configure_styles(doc)  # standard_business_brief preset
    base._configure_sections(doc)

    _cover(
        doc,
        kicker="阶段工作成果报告",
        title="严格小样本条件下的\n孪生光储备池CPI预测",
        subtitle="模型优化、结果对比与失败原因分析",
        rows=[
            ["汇报人", "陈中远"],
            ["专业", "电子信息工程"],
            ["个人分工", "孪生网络与仿真光储备池联动、孪生侧优化及公平对比"],
            ["报告周期", "2026年7月26日至2026年7月31日"],
        ],
        takeaway=(
            "严格封闭50样本下，SSA孪生模型测试MAE相对单光储备池下降2.43%，"
            "测试RMSE上升0.74%。孪生结构取得局部改善，但尚未形成全面稳定优势。"
        ),
    )

    base._add_heading(doc, "摘要", 1)
    base._add_body(
        doc,
        "本报告总结陈中远自2026年7月26日至7月31日在孪生光储备池小样本CPI预测"
        "方向完成的工作。研究首先建立严格封闭50样本协议，随后对最小时间间隔、"
        "参考距离、孪生特征、训练配对和聚合方式进行优化，并使用麻雀搜索算法SSA"
        "联合搜索孪生模型独有参数。在项目组统一环比数据、50/45/47时间划分和"
        "50维光储备池状态后，完成单光、孪生及传统模型的统一比较。"
    )
    base._add_body(
        doc,
        "正式环比实验中，SSA孪生模型测试MAE为0.303929，比单光储备池下降2.43%；"
        "测试RMSE为0.400142，比单光储备池上升0.74%。配对结构优化和SSA集成"
        "虽改善验证指标，但没有迁移到测试期。结果说明当前主要瓶颈已从流程实现"
        "转变为严格小样本和时间分布变化条件下的泛化能力。"
    )

    base._add_heading(doc, "1. 项目任务与个人分工", 1)
    base._add_body(
        doc,
        "项目使用软件仿真光储备池，不涉及实体光学硬件。杜艾佳负责单光储备池模型"
        "及统一状态生成，黄唯珂负责传统基线模型。陈中远负责孪生光储备池："
        "建立历史参考机制、优化孪生独有参数、执行公平性审计，并分析孪生结构"
        "相对单光储备池的有效性与局限。"
    )
    base._add_body(doc, "本阶段坚持两个边界：")
    for item in (
        "不改变光储备池随机掩码、虚拟节点、输入增益、延迟参数和50维状态；",
        "只优化孪生模型独有的参考选择、配对、特征、聚合和读出参数。",
    ):
        base._add_bullet(doc, item)

    base._add_heading(doc, "2. 上周日至今的工作时间线", 1)
    base._add_table(
        doc,
        ["时间", "主要工作", "阶段产出"],
        [
            [
                "7月26日",
                "建立严格封闭50样本；比较gap=1/3/6/12；测试混合距离和100维特征；完成单因素消融。",
                "明确100维特征和固定等权混合距离不适合当前小样本。",
            ],
            [
                "7月26日",
                "引入SSA，只优化gap、K、β、p、M等孪生独有参数。",
                "完成多随机种子自动搜索及收敛记录。",
            ],
            [
                "7月27—29日",
                "核对单光、孪生和传统基线的数据口径；确认旧30样本结果不能直接比较。",
                "确定统一环比序列、连续12个月窗口和固定时间划分。",
            ],
            [
                "7月30日",
                "接收并审计统一50维光储备池状态；重新运行正式环比SSA孪生实验和传统基线。",
                "形成50/45/47统一对比与正式核心结果。",
            ],
            [
                "7月31日",
                "尝试配对平衡、镜像样本、反对称回归、稳健聚合、SSA集成和残差修正。",
                "识别验证过拟合与时间分布变化问题，保留负向消融。",
            ],
            [
                "同期",
                "完善运行脚本、逐月预测表、实验清单、图表、数据泄漏检查和自动测试。",
                "35项仓库测试全部通过，结果可复现。",
            ],
        ],
        [1450, 4600, 3310],
    )

    doc.add_page_break()
    base._add_heading(doc, "3. 数据来源与实验公平性", 1)
    base._add_heading(doc, "3.1 数据来源", 2)
    base._add_body(
        doc,
        "最终数据来自 `data_processed/cpi_data_lastmonth=100.csv` 的 `actual` 列，"
        "含义为以上月为100的CPI环比指数。该序列与仓库原始环比数据一致，"
        "覆盖2001年8月至2026年5月且无缺失月份。正式实验不使用旧版预制特征。"
    )
    base._add_heading(doc, "3.2 样本构造和时间划分", 2)
    base._add_body(
        doc,
        "每个样本使用目标月份之前连续12个月的环比CPI，预测下一个月。"
        "所有模型按时间顺序使用相同目标月份："
    )
    base._add_table(
        doc,
        ["数据集", "目标月份", "样本数", "作用"],
        [
            ["训练集", "2014-07至2018-08", "50", "训练读出层和孪生差值模型"],
            ["验证集", "2018-09至2022-05", "45", "参数和方案选择"],
            ["测试集", "2022-06至2026-04", "47", "冻结配置后的评价"],
        ],
        [1800, 2800, 1200, 3560],
    )
    base._add_heading(doc, "3.3 光储备池状态", 2)
    base._add_body(
        doc,
        "杜艾佳使用MATLAB仿真生成统一状态，训练、验证、测试矩阵分别为50×50、"
        "45×50和47×50。陈中远核查了状态尺寸、样本编号、月份、目标值、连续窗口、"
        "共享随机掩码及有限性。单光和孪生模型使用完全相同的状态。"
    )
    base._add_heading(doc, "3.4 严格封闭50样本", 2)
    base._add_body(
        doc,
        "孪生模型所有训练对、验证参考和测试参考均限制在固定训练50窗口内。"
        "组合后训练对行数可以超过50，但没有引入新月份，保证与单光储备池"
        "具有相同原始数据访问预算。"
    )

    base._add_heading(doc, "4. 孪生光储备池方法", 1)
    base._add_body(
        doc,
        "单光储备池直接学习50维状态h_i到目标y_i的映射。孪生模型选择历史参考j，"
        "使用状态差h_i-h_j预测Δy_ij=y_i-y_j，再通过ŷ_i=y_j+Δŷ_ij还原当前目标。"
        "一个目标使用多个参考，最后根据参考距离聚合。"
    )
    base._add_callout(
        doc,
        "方法差异：",
        "单光预测绝对值；孪生模型参考历史相似窗口，预测当前相对历史变化了多少。",
    )

    doc.add_page_break()
    base._add_heading(doc, "5. 优化方法、选择理由和处理决定", 1)
    base._add_table(
        doc,
        ["方法", "选择理由", "主要结果", "最终决定"],
        [
            ["gap=1/3/6/12", "增加或约束合法配对", "缩小gap产生高度重叠配对，单因素验证未改善", "不单独缩小"],
            ["形状+水平距离", "避免只按形状选到水平不匹配的参考", "固定各50%时测试RMSE约上升5.36%", "放弃固定等权"],
            ["100维[h_i,h_i-h_j]", "保留目标状态与相对差异", "测试RMSE约上升52.73%，明显过拟合", "放弃，保留50维差"],
            ["SSA联合优化", "处理离散/连续参数及其相互作用", "正式环比验证MAE/RMSE降至0.353049/0.446791", "作为主优化方法"],
            ["配对结构优化", "改善目标和差值区间不均衡", "验证RMSE降至0.406930，测试升至0.422209", "负向消融"],
            ["SSA三成员集成", "降低单次搜索随机性", "验证小幅改善，测试RMSE为0.404185", "不替换主模型"],
            ["残差式修正", "用单光提供稳定基础、孪生负责修正", "验证选择λ=1，无额外修正价值", "不采用"],
        ],
        [1950, 3000, 2830, 1580],
    )
    base._add_callout(
        doc,
        "口径说明：",
        "gap、混合距离和100维特征的初始消融属于正式数据统一前的方法探索，"
        "用于决定后续结构；最终性能结论只使用统一后的环比50/45/47实验。",
        fill=base.PALE_GOLD,
        accent="7A5A00",
    )

    base._add_heading(doc, "6. SSA优化方案", 1)
    base._add_body(
        doc,
        "SSA能够同时处理gap、K、M等离散参数与β、p等连续参数，适合联合搜索。"
        "三个随机种子共评价346组不重复配置。目标函数为验证RMSE加三个时间块"
        "RMSE标准差的0.1倍，以兼顾整体误差与时间稳定性。"
    )
    base._add_table(
        doc,
        ["参数", "最终值", "作用与解释"],
        [
            ["gap", "1", "与K、M及强正则化共同选择"],
            ["K", "5", "减少低质量参考干扰"],
            ["β", "0.041", "参考距离以形状为主，水平为辅"],
            ["p", "0.654", "控制距离加权的集中程度"],
            ["M", "1", "每个差值区间保留少量代表性配对"],
            ["Ridge alpha", "100", "小样本下使用较强正则化"],
        ],
        [1900, 1500, 5960],
    )
    _add_report_figure(
        doc,
        base.FIG_SSA,
        "图1  三个随机种子的SSA验证集收敛过程",
        "SSA在三个随机种子下的验证目标收敛曲线。",
        6.0,
    )

    doc.add_page_break()
    base._add_heading(doc, "7. 正式环比实验结果", 1)
    base._add_heading(doc, "7.1 单光与孪生核心对比", 2)
    base._add_table(
        doc,
        ["模型", "验证MAE", "验证RMSE", "测试MAE", "测试RMSE"],
        [
            ["单光储备池", "0.386178", "0.487698", "0.311511", "0.397189"],
            ["未优化孪生", "0.370512", "0.467292", "0.319792", "0.406210"],
            ["SSA优化孪生", "0.353049", "0.446791", "0.303929", "0.400142"],
        ],
        [2500, 1715, 1715, 1715, 1715],
    )
    base._add_body(
        doc,
        "SSA孪生测试MAE相对单光下降2.43%，说明多数月份平均绝对误差降低；"
        "测试RMSE相对单光上升0.74%，说明少数较大误差仍未得到稳定控制。"
    )
    _add_report_figure(
        doc,
        base.FIG_KEY_PRED,
        "图2  47个测试月份的真实值、单光预测与孪生预测",
        "47个测试月份的CPI真实值和关键模型预测曲线。",
        6.25,
    )

    base._add_heading(doc, "7.2 传统模型补充对照", 2)
    base._add_table(
        doc,
        ["模型", "测试MAE", "测试RMSE"],
        [
            ["Gradient Boosting", "0.305704", "0.382399"],
            ["Ridge", "0.296734", "0.395175"],
            ["单光储备池", "0.311511", "0.397189"],
            ["SSA孪生光储备池", "0.303929", "0.400142"],
            ["随机森林", "0.327827", "0.402133"],
            ["未优化孪生", "0.319792", "0.406210"],
        ],
        [4360, 2500, 2500],
    )
    _add_report_figure(
        doc,
        base.FIG_ALL_MODELS,
        "图3  相同数据与时间划分下全部模型的测试指标",
        "传统模型、单光储备池和孪生光储备池的测试MAE与RMSE对比。",
        6.15,
    )

    doc.add_page_break()
    base._add_heading(doc, "8. 后续探索及未采用原因", 1)
    base._add_heading(doc, "8.1 配对结构优化", 2)
    base._add_body(
        doc,
        "进一步测试全部合法配对、目标/差值区间平衡权重、镜像样本、零截距反对称"
        "回归和截尾均值聚合。验证RMSE降至0.406930，但测试RMSE升至0.422209，"
        "比单光高6.30%，因此不采用。"
    )
    _add_report_figure(
        doc,
        base.FIG_PAIR,
        "图4  配对结构候选方案的验证集表现",
        "不同gap与训练模式在验证集上的表现。",
        6.1,
    )
    base._add_heading(doc, "8.2 SSA集成与残差修正", 2)
    base._add_body(
        doc,
        "三成员等权集成将验证RMSE小幅降至0.444837，但测试RMSE为0.404185。"
        "三个成员共享数据、状态和结构，误差相关性高，集成降低方差的作用有限。"
        "残差式修正的验证最优λ=1，说明没有从单光预测获得额外修正价值。"
    )
    _add_report_figure(
        doc,
        base.FIG_ENSEMBLE,
        "图5  单光、最佳SSA、SSA集成及残差修正的测试指标",
        "后续孪生优化方案与单光储备池的测试误差对比。",
        6.1,
    )

    base._add_heading(doc, "9. 优化失败原因分析", 1)
    base._add_table(
        doc,
        ["原因", "影响"],
        [
            ["原始训练窗口仅50个", "难以支持高维特征和复杂训练结构"],
            ["配对并非独立样本", "组合增加行数，但没有增加新的经济时期"],
            ["滑动窗口高度重叠", "相邻窗口有效信息增量有限"],
            ["训练与测试阶段不同", "2014—2018参考池难覆盖2022—2026经济状态"],
            ["固定验证区间单一", "SSA和结构优化可能适配45个月验证期"],
            ["模型成员误差相关", "共享数据、状态和结构，集成收益有限"],
            ["环比信号接近100", "短期波动和噪声相对有效信号占比较高"],
        ],
        [2800, 6560],
    )
    base._add_callout(
        doc,
        "综合判断：",
        "当前瓶颈不是联动流程没有跑通，而是严格小样本下的跨时间阶段泛化。",
        fill=base.PALE_GOLD,
        accent="7A5A00",
    )

    doc.add_page_break()
    base._add_heading(doc, "10. 其他完成工作", 1)
    for item in (
        "核对数据口径，明确最终任务为直接预测CPI环比；",
        "统一单光、孪生和传统模型的训练、验证、测试月份；",
        "审计MATLAB状态与原始目标、月份、样本编号和共享掩码；",
        "建立严格封闭参考池和防数据泄漏检查；",
        "补充MAT状态读取兼容，完善可复现运行脚本；",
        "输出逐月预测、指标表、收敛曲线、模型清单和实验边界说明；",
        "完成传统模型统一对照；",
        "仓库完整自动测试35项全部通过。",
    ):
        base._add_bullet(doc, item)

    base._add_heading(doc, "11. 阶段成果与结论", 1)
    base._add_body(doc, "本阶段形成的主要成果包括：")
    for item in (
        "完整跑通仿真光储备池与孪生网络联动预测流程；",
        "建立严格封闭50样本公平比较协议；",
        "完成多种孪生结构的单因素消融和自动优化；",
        "形成正式环比50/45/47统一对比结果；",
        "明确100维特征、固定等权混合距离、配对强化和模型集成未形成稳定收益；",
        "识别验证过拟合与时间分布变化为主要瓶颈。",
    ):
        base._add_bullet(doc, item)
    base._add_callout(
        doc,
        "阶段结论：",
        "孪生历史参考机制在测试MAE上带来2.43%的改善，证明其具有一定价值；"
        "但测试RMSE仍高0.74%，尚不能认为孪生光储备池全面优于单光储备池。",
    )
    base._add_body(
        doc,
        "后续建议冻结当前主方案，停止根据同一测试区间继续调参，"
        "采用多个滚动时间验证区间，并使用2026年5月之后的新月份作为真正未见数据。"
    )

    base._add_heading(doc, "附录：关键数据文件", 1)
    base._add_table(
        doc,
        ["内容", "仓库位置"],
        [
            ["原始环比数据", "data_processed/cpi_data_lastmonth=100.csv"],
            ["统一50维状态", "matlab/optical_reservoir_cpi_mom_recent50_20260730/states"],
            ["环比统一比较", "results/mom_closed50_unified_comparison_20260730"],
            ["正式SSA孪生实验", "results/siamese_optical_mom_closed50_ssa_20260730"],
            ["配对结构探索", "results/siamese_optical_mom_pair_optimization_20260731"],
            ["集成与残差探索", "results/siamese_optical_mom_ssa_ensemble_residual_20260731"],
        ],
        [2600, 6760],
    )

    doc.core_properties.title = "孪生光储备池阶段工作成果报告"
    doc.core_properties.subject = "2026年7月26日至7月31日阶段工作"
    doc.core_properties.author = "陈中远"
    doc.core_properties.keywords = "CPI, 孪生光储备池, SSA, 小样本"
    for section in doc.sections:
        _set_running_furniture(section, "孪生光储备池阶段工作成果报告｜陈中远")
    REPORT_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_OUTPUT)
    return REPORT_OUTPUT


def _add_speech(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.keep_together = False
    run = paragraph.add_run(text)
    base._set_run_font(run, size=11, color=base.INK)


def _add_stage_direction(doc: Document, text: str) -> None:
    base._add_callout(
        doc,
        "展示提示：",
        text,
        fill=base.PALE_GOLD,
        accent="7A5A00",
    )


def build_speaking_script() -> Path:
    doc = Document()
    base._configure_styles(doc)
    _apply_compact_guide_styles(doc)  # compact_reference_guide preset
    base._configure_sections(doc)

    _cover(
        doc,
        kicker="配套讲解稿",
        title="孪生光储备池\n阶段工作成果汇报",
        subtitle="按正式报告章节和图表顺序编排",
        rows=[
            ["汇报人", "陈中远"],
            ["配套报告", "《孪生光储备池阶段工作成果报告》"],
            ["建议时长", "8—12分钟"],
            ["使用方式", "正文可直接照读；黄色提示不需要读出"],
        ],
        takeaway=(
            "先讲公平性和工作过程，再讲SSA核心结果；失败方法重点解释选择理由和"
            "放弃原因，最后用“MAE局部改善、RMSE尚未稳定”收束。"
        ),
    )

    base._add_heading(doc, "汇报节奏", 1)
    base._add_table(
        doc,
        ["部分", "建议时间", "对应报告"],
        [
            ["开场与任务", "1分钟", "第1—2节"],
            ["数据来源与公平性", "2分钟", "第3—4节"],
            ["优化方法与选择理由", "3分钟", "第5—6节"],
            ["核心结果", "2分钟", "第7节、图1—3"],
            ["后续探索与失败原因", "2分钟", "第8—9节、图4—5"],
            ["其他工作与结论", "1—2分钟", "第10—11节"],
        ],
        [2700, 1900, 4760],
    )

    base._add_heading(doc, "1. 开场", 1)
    _add_speech(
        doc,
        "老师好，我是陈中远。我在项目中负责孪生光储备池部分，主要任务是把孪生网络"
        "和软件仿真的光储备池结合起来，在小样本条件下预测CPI，并与单光储备池进行"
        "公平比较。"
    )
    _add_speech(
        doc,
        "从上周日，也就是7月26日开始到现在，我主要完成了四类工作：首先建立严格"
        "封闭50样本实验；然后测试gap、参考距离和100维特征；接着使用SSA联合优化"
        "孪生独有参数；最后在统一环比数据以后，继续尝试配对结构、模型集成和残差修正。"
    )
    _add_speech(
        doc,
        "我先给出最终结论：SSA孪生模型测试MAE相对单光储备池下降了2.43%，但测试"
        "RMSE上升了0.74%。所以目前取得的是局部改善，还不能说孪生模型全面优于单光。"
    )

    base._add_heading(doc, "2. 数据来源与公平性", 1)
    _add_stage_direction(doc, "打开正式报告第3节，展示50/45/47数据划分表。")
    _add_speech(
        doc,
        "最终数据来自仓库cpi_data_lastmonth=100.csv中的actual列。这是以上个月为100"
        "的CPI环比指数。我们只使用这一列，重新按照时间顺序构造连续过去12个月预测"
        "下一个月的滑动窗口，不使用旧版预制特征。"
    )
    _add_speech(
        doc,
        "训练目标是2014年7月至2018年8月，共50个；验证是2018年9月至2022年5月，"
        "共45个；测试是2022年6月至2026年4月，共47个。这个划分完全按照时间顺序，"
        "没有随机打乱。"
    )
    _add_speech(
        doc,
        "杜艾佳使用MATLAB仿真生成统一的50维光储备池状态。我没有另生成一套状态，"
        "而是检查状态尺寸、样本编号、月份、目标值和共享掩码，确认单光和孪生使用"
        "完全相同的50维状态。"
    )
    _add_speech(
        doc,
        "我还建立了严格封闭50样本规则。孪生的训练对和验证、测试参考都只能来自固定"
        "训练50窗口。虽然50个窗口可以组合出多行训练对，但没有增加新的月份，所以与"
        "单光模型的数据预算一致。"
    )

    base._add_heading(doc, "3. 孪生模型原理", 1)
    _add_speech(
        doc,
        "单光储备池直接根据目标窗口的50维状态预测下一个月CPI。孪生模型会给当前窗口"
        "找一个历史参考窗口，用两个状态之差预测两个目标值之差，再把预测差值加到参考"
        "窗口的真实CPI上。一个目标使用多个参考，最后按照参考距离聚合。"
    )
    _add_speech(
        doc,
        "所以单光模型学习的是状态到绝对值，孪生模型学习的是当前相对历史变化了多少。"
    )

    base._add_heading(doc, "4. 上周日首先尝试的优化", 1)
    _add_stage_direction(doc, "翻到正式报告第5节的“优化方法、选择理由和处理决定”表。")
    _add_speech(
        doc,
        "第一项是调整gap。我比较了1、3、6、12，希望缩小gap以后产生更多训练对。"
        "但相邻12个月窗口可能有11个月相同，所以增加的大多是高度相关配对。单因素验证"
        "没有改善，因此没有采用单纯缩小gap的方法。"
    )
    _add_speech(
        doc,
        "第二项是混合参考距离。我让参考选择同时考虑变化形状和绝对水平，希望避免只按"
        "形状选到水平差异过大的参考。但固定各占一半以后，测试RMSE约上升5.36%。"
        "我推测绝对水平接近不一定代表后续环比规律接近，所以放弃固定等权，只让SSA保留"
        "很小的水平权重。"
    )
    _add_speech(
        doc,
        "第三项是100维特征。我把原来的50维状态差扩展为目标状态加状态差，希望同时"
        "保留目标本身和相对信息。但测试RMSE约上升52.73%，退化最明显。原因是原始"
        "训练窗口只有50个，配对也不是独立新样本，100维特征造成了明显过拟合。"
        "因此最终保留50维状态差。"
    )

    base._add_heading(doc, "5. 为什么选择SSA", 1)
    _add_stage_direction(doc, "展示报告图1：SSA三个随机种子的验证收敛曲线。")
    _add_speech(
        doc,
        "完成单因素实验以后，我发现gap、参考数K、水平权重、聚合指数和配对数之间"
        "存在相互作用，所以不能把单因素最优值简单组合。"
    )
    _add_speech(
        doc,
        "我选择SSA，是因为它可以同时搜索离散和连续参数，而且只需要优化孪生模型独有"
        "参数，不改变共享的光储备池。三个随机种子一共评价346组不重复配置。"
    )
    _add_speech(
        doc,
        "最终选择gap等于1、K等于5、水平权重0.041、距离指数0.654、每个差值区间"
        "最多1对，Ridge正则参数100。这个结果说明参考选择仍以形状为主，而且小样本下"
        "需要少量参考和较强正则化。"
    )

    base._add_heading(doc, "6. 核心结果", 1)
    _add_stage_direction(doc, "展示正式报告第7.1节核心指标表，然后展示图2预测曲线。")
    _add_speech(
        doc,
        "单光储备池测试MAE和RMSE分别是0.311511和0.397189。未优化孪生是0.319792"
        "和0.406210。SSA优化孪生是0.303929和0.400142。"
    )
    _add_speech(
        doc,
        "相对单光，SSA孪生的测试MAE下降2.43%，说明多数月份平均误差减小；但RMSE"
        "上升0.74%，说明少数大误差月份没有控制好。因此我的表述是局部改善，而不是"
        "全面优于单光。"
    )
    _add_stage_direction(doc, "如老师关心传统模型，再展示报告图3；否则可略过。")
    _add_speech(
        doc,
        "传统模型只是补充参照。同一时间划分下，Gradient Boosting的RMSE最低，"
        "Ridge的MAE最低。这也说明50样本条件下，强正则化传统模型仍然很有竞争力。"
    )

    base._add_heading(doc, "7. 后续方法为什么放弃", 1)
    _add_stage_direction(doc, "展示报告图4和图5，重点讲验证提高、测试退化。")
    _add_speech(
        doc,
        "随后我从配对结构继续优化，加入全部合法配对、平衡权重、镜像样本、反对称回归"
        "和截尾均值。验证RMSE从0.446791降到0.406930，但测试RMSE升到0.422209，"
        "比单光高6.30%。因此没有采用。"
    )
    _add_speech(
        doc,
        "我认为原因是配对数量虽然增加，但仍来自同样50个窗口，相关性很强；同时模型"
        "适应了2018到2022年的验证期，却没有适应2022到2026年的测试期。"
    )
    _add_speech(
        doc,
        "我又把三个SSA种子的最佳模型等权集成。验证有小幅改善，但测试RMSE为0.404185。"
        "三个成员共享相同数据、状态和结构，错误高度相关，所以集成降低方差的作用有限。"
    )
    _add_speech(
        doc,
        "最后尝试单光加孪生残差修正，验证集选择的权重是1，也就是完全采用孪生集成，"
        "单光没有提供额外修正价值，所以这一方案也没有采用。"
    )

    base._add_heading(doc, "8. 除了调模型还完成了什么", 1)
    _add_speech(
        doc,
        "除了模型优化，我还统一了环比数据口径，核对了50、45、47的时间划分，审计了"
        "MATLAB状态，建立了防数据泄漏检查，补充了MAT状态读取兼容，重新运行传统模型，"
        "并输出逐月预测、参数表、收敛图和实验清单。目前仓库35项自动测试全部通过。"
    )

    base._add_heading(doc, "9. 结束总结", 1)
    _add_speech(
        doc,
        "从上周日到现在，我完成的不是单次参数搜索，而是从公平性定义、结构消融、SSA"
        "优化到失败分析的完整过程。最终保留严格封闭50样本、共享50维状态、50维状态差、"
        "以形状为主的参考距离和SSA选择的孪生参数。"
    )
    _add_speech(
        doc,
        "当前孪生模型在MAE上证明了历史参考机制具有一定价值，但RMSE还没有形成稳定优势。"
        "我认为下一步不应该继续根据同一个测试集调参，而应该冻结当前方案，使用滚动时间"
        "验证，并等待新的未见月份评价。我的汇报结束，谢谢老师。"
    )

    base._add_heading(doc, "10. 常见追问速答", 1)
    qa = [
        ["为什么配对数超过50仍公平？", "所有配对都来自固定50个窗口，没有访问额外月份。"],
        ["孪生模型最终是否更好？", "MAE好2.43%，RMSE差0.74%，属于局部改善。"],
        ["为什么不采用验证更好的配对模型？", "测试明显退化，缺乏跨时间稳定性。"],
        ["为什么100维失败？", "小样本、高维、配对相关，导致过拟合。"],
        ["为什么集成没有效果？", "成员共享数据、状态和结构，错误相关性高。"],
        ["下一步是什么？", "冻结方案，滚动验证，并使用新的未来月份。"],
    ]
    base._add_table(doc, ["问题", "回答"], qa, [3400, 5960])

    doc.core_properties.title = "孪生光储备池阶段工作成果报告讲解稿"
    doc.core_properties.subject = "配套口头汇报稿"
    doc.core_properties.author = "陈中远"
    for section in doc.sections:
        _set_running_furniture(section, "阶段工作成果报告讲解稿｜陈中远")
    SCRIPT_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SCRIPT_OUTPUT)
    return SCRIPT_OUTPUT


if __name__ == "__main__":
    print(build_report())
    print(build_speaking_script())
