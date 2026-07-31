"""Build Chen Zhongyuan's final oral project report as a polished DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "陈中远_孪生光储备池最终口头汇报报告_20260731.docx"

FIG_KEY_PRED = (
    ROOT
    / "results"
    / "mom_closed50_unified_comparison_20260730"
    / "figures"
    / "key_model_test_predictions.png"
)
FIG_ALL_MODELS = (
    ROOT
    / "results"
    / "mom_closed50_unified_comparison_20260730"
    / "figures"
    / "all_model_test_metrics.png"
)
FIG_SSA = (
    ROOT
    / "results"
    / "siamese_optical_mom_closed50_ssa_20260730"
    / "figures"
    / "ssa_validation_convergence.png"
)
FIG_PAIR = (
    ROOT
    / "results"
    / "siamese_optical_mom_pair_optimization_20260731"
    / "figures"
    / "validation_pair_structure_comparison.png"
)
FIG_ENSEMBLE = (
    ROOT
    / "results"
    / "siamese_optical_mom_ssa_ensemble_residual_20260731"
    / "figures"
    / "test_metric_comparison.png"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203040"
MUTED = "626B75"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF5D8"
GREEN = "2F6B4F"
RED = "9B1C1C"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    east_asia: str = "Microsoft YaHei",
    latin: str = "Calibri",
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = "CCD3DA", size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths sum to {sum(widths_dxa)}, expected {TABLE_WIDTH_DXA}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _style_table(table, *, first_col_left: bool = True) -> None:
    _set_repeat_table_header(table.rows[0])
    _set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            for cell in row.cells:
                _set_cell_shading(cell, BLUE)
        elif row_index % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, LIGHT_GRAY)
        for col_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if first_col_left and col_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    _set_run_font(
                        run,
                        size=9.0,
                        bold=row_index == 0,
                        color=WHITE if row_index == 0 else INK,
                    )


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    first_col_left: bool = True,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    _set_table_geometry(table, widths_dxa)
    _style_table(table, first_col_left=first_col_left)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_body(
    doc: Document,
    text: str,
    *,
    bold_prefix: str | None = None,
    italic: bool = False,
    color: str = INK,
    keep_with_next: bool = False,
) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        _set_run_font(first, size=11, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        _set_run_font(rest, size=11, italic=italic, color=color)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run, size=11, italic=italic, color=color)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    _set_run_font(run, size=10.5, color=INK)


def _add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    _set_run_font(run, size=10.5, color=INK)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    _set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color=BLUE if level < 3 else DARK_BLUE,
    )


def _add_callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = LIGHT_BLUE,
    accent: str = DARK_BLUE,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_table_geometry(table, [TABLE_WIDTH_DXA])
    _set_table_borders(table, color=fill, size=2)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    lead = p.add_run(f"{label} ")
    _set_run_font(lead, size=10.5, bold=True, color=accent)
    body = p.add_run(text)
    _set_run_font(body, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_figure(
    doc: Document,
    image_path: Path,
    caption: str,
    *,
    width: float = 6.25,
    oral_prompt: str | None = None,
) -> None:
    if oral_prompt:
        _add_callout(doc, "展示提示：", oral_prompt, fill=PALE_GOLD, accent="7A5A00")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(2)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    _set_run_font(caption_run, size=9, italic=True, color=MUTED)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(end)
    tail = paragraph.add_run(" 页")
    _set_run_font(tail, size=9, color=MUTED)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def _configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True


def _add_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("孪生光储备池小样本CPI预测｜最终汇报")
    _set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    _add_page_number(p)


def _add_cover(doc: Document) -> None:
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run("项目阶段总结与最终口头汇报")
    _set_run_font(kicker_run, size=12, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("严格小样本条件下的\n孪生光储备池CPI预测")
    _set_run_font(title_run, size=27, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    subtitle_run = subtitle.add_run("方法优化、失败分析与阶段结论")
    _set_run_font(subtitle_run, size=15, color=MUTED)

    _add_table(
        doc,
        ["项目", "内容"],
        [
            ["汇报人", "陈中远"],
            ["专业", "电子信息工程"],
            ["个人分工", "孪生网络与仿真光储备池联动、孪生模型优化及公平对比"],
            ["总结周期", "2026年7月26日至2026年7月31日"],
        ],
        [2100, 7260],
        first_col_left=False,
    )

    _add_callout(
        doc,
        "核心结果：",
        "SSA优化孪生模型测试MAE相对单光储备池下降2.43%，"
        "测试RMSE上升0.74%。孪生结构取得局部改善，但尚未全面优于单光储备池。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )
    doc.add_page_break()


def build_document() -> Path:
    for image in (FIG_KEY_PRED, FIG_ALL_MODELS, FIG_SSA, FIG_PAIR, FIG_ENSEMBLE):
        if not image.exists():
            raise FileNotFoundError(image)

    doc = Document()
    _configure_styles(doc)
    _configure_sections(doc)
    _add_cover(doc)

    _add_heading(doc, "使用说明与汇报重点", 1)
    _add_body(
        doc,
        "本报告正文采用口头化表述，可直接照读或根据现场时间删减。"
        "黄色“展示提示”用于提醒何时切换图表，不需要读给老师听。"
    )
    _add_callout(
        doc,
        "建议开场结论：",
        "我从上周日开始，围绕严格封闭50样本下的孪生光储备池优化，"
        "先后尝试gap、混合距离、100维特征、SSA、配对结构、模型集成和残差修正。"
        "最终SSA模型改善了MAE，但RMSE没有形成稳定优势。",
    )
    _add_body(doc, "建议把汇报重点放在以下四件事上：")
    for item in (
        "为什么必须采用严格封闭50样本，保证与单光储备池公平比较；",
        "为什么选择这些优化方法，以及每种方法解决什么问题；",
        "为什么部分方法验证集提高、测试集却退化；",
        "最终保留什么方案，以及能够得到什么程度的结论。",
    ):
        _add_bullet(doc, item)

    _add_heading(doc, "一、开场：我的任务与本周工作主线", 1)
    _add_body(
        doc,
        "老师好，我是陈中远。我在项目中负责孪生光储备池部分，"
        "主要任务是把孪生网络和软件仿真的光储备池结合起来，"
        "在小样本条件下预测CPI，并与杜艾佳负责的单光储备池进行公平比较。"
    )
    _add_body(
        doc,
        "从上周日，也就是7月26日开始到现在，我主要完成了四个阶段的工作："
    )
    for item in (
        "建立严格封闭50样本实验，并对gap、参考距离和孪生特征进行单因素分析；",
        "使用麻雀搜索算法SSA联合优化孪生模型独有参数；",
        "与组内统一正式环比数据、时间划分和50维光储备池状态，重新运行公平对比；",
        "继续尝试配对结构优化、SSA多模型集成和残差式修正，并分析失败原因。",
    ):
        _add_numbered(doc, item)
    _add_body(
        doc,
        "我先汇报最终结论：当前孪生光储备池已经完整跑通。"
        "SSA优化以后，测试MAE比单光储备池下降2.43%，但测试RMSE上升0.74%。"
        "因此只能说孪生结构取得了局部改善，还不能说全面优于单光储备池。"
    )

    doc.add_page_break()
    _add_heading(doc, "二、数据从哪里来", 1)
    _add_heading(doc, "2.1 原始数据及目标定义", 2)
    _add_body(
        doc,
        "我们最终统一使用仓库 `data_processed/cpi_data_lastmonth=100.csv` "
        "中的 `actual` 列。这是以上个月为100的CPI环比指数。"
        "例如100.5表示相对上个月上涨约0.5%，99.8表示相对上个月下降约0.2%。"
    )
    _add_body(
        doc,
        "我核对过这列数据，它与仓库中的原始CPI环比序列一致，"
        "时间从2001年8月延续到2026年5月，中间没有缺失月份。"
        "正式实验不再使用旧版 `month-19`、`month-17` 等预制特征。"
    )
    _add_heading(doc, "2.2 滑动窗口如何构造", 2)
    _add_body(
        doc,
        "每个样本使用目标月份之前连续12个月的环比CPI作为输入，"
        "预测下一个月的环比CPI。比如预测2018年9月，输入就是2017年9月至2018年8月。"
    )
    _add_table(
        doc,
        ["数据集", "目标月份", "目标数", "用途"],
        [
            ["训练集", "2014-07至2018-08", "50", "训练单光和孪生读出层"],
            ["验证集", "2018-09至2022-05", "45", "选择参数和模型配置"],
            ["测试集", "2022-06至2026-04", "47", "配置冻结后的最终评价"],
        ],
        [1800, 2800, 1200, 3560],
    )
    _add_callout(
        doc,
        "为什么测试集是47个：",
        "测试目标按项目组统一的固定日历区间2022年6月至2026年4月确定，"
        "该区间正好包含47个月。样本不是随机打乱后按比例截取的。",
        fill=PALE_GOLD,
        accent="7A5A00",
    )

    _add_heading(doc, "2.3 光储备池状态如何获得", 2)
    _add_body(
        doc,
        "项目没有实体光储备池硬件，使用软件仿真光储备池。"
        "每个12个月输入窗口经过随机掩码、虚拟节点和非线性映射，得到一个50维状态。"
    )
    _add_table(
        doc,
        ["数据集", "状态矩阵", "生成与使用"],
        [
            ["训练集", "50×50", "50个窗口，每个窗口50维"],
            ["验证集", "45×50", "45个窗口，每个窗口50维"],
            ["测试集", "47×50", "47个窗口，每个窗口50维"],
        ],
        [2100, 2000, 5260],
    )
    _add_body(
        doc,
        "杜艾佳使用MATLAB按照统一数据生成这套状态。我的电脑没有安装MATLAB，"
        "所以我没有另生成一套状态，而是复用完全相同的状态，保证单光与孪生比较公平。"
    )
    _add_body(doc, "收到状态以后，我完成了以下一致性检查：")
    for item in (
        "状态尺寸、样本编号、月份和目标值是否一一对应；",
        "输入是否确实为目标月份之前连续12个月；",
        "训练、验证和测试是否共享同一个随机掩码；",
        "是否存在缺失值、无穷值或全常数状态；",
        "单光和孪生是否使用完全相同的50维状态。",
    ):
        _add_bullet(doc, item)
    _add_body(doc, "上述检查全部通过。")

    doc.add_page_break()
    _add_heading(doc, "三、单光与孪生模型的区别", 1)
    _add_heading(doc, "3.1 单光储备池", 2)
    _add_body(
        doc,
        "单光储备池直接使用目标窗口的50维状态 h_i，通过Ridge回归预测目标CPI y_i。"
        "它学习的是“状态到CPI绝对值”的映射。"
    )
    _add_callout(doc, "单光预测：", "h_i  →  y_i")
    _add_heading(doc, "3.2 孪生光储备池", 2)
    _add_body(
        doc,
        "孪生模型同时使用目标窗口i和历史参考窗口j。"
        "两个窗口通过同一个冻结的光储备池得到状态h_i和h_j，"
        "模型使用状态差h_i-h_j预测目标差值Δy_ij=y_i-y_j。"
    )
    _add_callout(
        doc,
        "孪生预测：",
        "先预测 Δy_ij，再用参考窗口真实值还原：ŷ_i = y_j + Δŷ_ij。",
    )
    _add_body(
        doc,
        "每个目标选择多个历史参考窗口，每个参考产生一个候选预测，"
        "最后按照参考距离加权聚合。孪生网络的作用不是替代光储备池，"
        "而是把直接预测转化为“参考历史相似状态，预测当前相对历史变化了多少”。"
    )

    _add_heading(doc, "四、为什么建立严格封闭50样本实验", 1)
    _add_body(
        doc,
        "早期孪生模型虽然只有50个监督训练目标，但还可以访问额外历史参考窗口，"
        "实际数据预算超过单光储备池。因此我重新规定："
    )
    for item in (
        "整个模型只能访问固定的50个训练窗口；",
        "训练样本对只能由这50个窗口组合产生；",
        "验证和测试参考窗口也只能来自训练50；",
        "验证和测试真实目标不能进入训练参考池；",
        "参数只能使用验证集选择，测试集仅在配置冻结后评价。",
    ):
        _add_bullet(doc, item)
    _add_body(
        doc,
        "孪生模型可以由50个窗口组合出多个训练样本对，但这些配对没有引入新月份，"
        "只是重新组织已有信息。因此两种模型的原始数据访问预算保持一致。"
    )
    _add_callout(
        doc,
        "汇报重点：",
        "公平性比较的是模型可以访问多少个原始月份，不是孪生组合以后产生多少行样本对。",
        fill=PALE_GOLD,
        accent="7A5A00",
    )

    doc.add_page_break()
    _add_heading(doc, "五、7月26日：三项结构优化及放弃原因", 1)
    _add_heading(doc, "5.1 调整最小时间间隔gap", 2)
    _add_body(
        doc,
        "我比较了gap=1、3、6、12。选择gap进行优化，是因为它决定目标窗口和参考窗口"
        "之间至少间隔多长时间，也会影响可生成的训练样本对数量。"
    )
    _add_body(
        doc,
        "缩小gap虽然产生更多配对，但相邻12个月窗口可能有11个月完全相同。"
        "单因素实验中，验证集仍选择gap=12，说明增加的是高度相关的重复配对，"
        "有效信息没有同步增加。因此没有采用“单纯减小gap”的方案。"
    )

    _add_heading(doc, "5.2 同时考虑形状与绝对水平", 2)
    _add_body(
        doc,
        "我尝试让参考距离同时考虑12个月变化形状、窗口均值和窗口最后一个月水平。"
        "选择这个方向，是因为只看形状可能选到趋势相似但经济水平不同的参考。"
    )
    _add_body(
        doc,
        "固定采用形状和水平各50%以后，测试RMSE相对原孪生基线上升约5.36%。"
        "我推测绝对水平接近并不保证下一阶段的环比变化规律接近，"
        "反而可能排除形状更合适的参考。因此放弃固定等权混合，"
        "只保留由SSA自动选择的少量水平权重。"
    )

    _add_heading(doc, "5.3 使用100维特征", 2)
    _add_body(
        doc,
        "原特征是50维状态差h_i-h_j。我尝试拼接目标状态和状态差，"
        "形成100维特征[h_i, h_i-h_j]，希望同时保留目标本身及相对参考的信息。"
    )
    _add_body(
        doc,
        "该方案测试RMSE相对原孪生基线上升约52.73%，是退化最明显的方法。"
        "主要原因是只有50个原始训练窗口，组合样本对之间也高度相关，"
        "100维特征带来明显的小样本高维过拟合。因此最终继续使用50维状态差。"
    )
    _add_table(
        doc,
        ["尝试方法", "选择理由", "结果", "处理"],
        [
            ["缩小gap", "增加合法配对", "高度重叠，验证未改善", "不单独采用"],
            ["混合距离", "兼顾形状和水平", "测试RMSE约+5.36%", "放弃固定等权"],
            ["100维特征", "保留目标状态信息", "测试RMSE约+52.73%", "放弃，保留50维"],
        ],
        [1750, 2700, 2600, 2310],
    )
    _add_callout(
        doc,
        "口头说明：",
        "这组三项消融属于正式数据口径统一前的方法探索，主要用于决定后续模型结构；"
        "最终性能结论只采用后面统一环比数据重新运行的结果。",
        fill=PALE_GOLD,
        accent="7A5A00",
    )

    doc.add_page_break()
    _add_heading(doc, "六、为什么选择SSA联合优化", 1)
    _add_body(
        doc,
        "单因素实验以后，我发现孪生参数存在相互作用。gap影响合法配对数量，"
        "K影响参考稳定性，β影响参考选择，p影响聚合权重，M影响训练分布。"
        "因此不能简单把每个单因素最优值直接拼在一起。"
    )
    _add_body(doc, "我选择麻雀搜索算法SSA，主要有三个原因：")
    for item in (
        "可以同时搜索gap、K等离散参数和β、p等连续参数；",
        "适合处理多个参数之间的联合影响，比逐项手工调整更系统；",
        "只优化孪生模型独有参数，不改变共享光储备池，符合个人分工和公平性要求。",
    ):
        _add_numbered(doc, item)
    _add_body(
        doc,
        "正式环比实验使用三个随机种子，共评价346组不重复配置。"
        "目标函数不仅包含验证RMSE，还加入验证集三个时间块RMSE标准差的惩罚，"
        "避免模型只在验证区间中的某一段表现较好。"
    )
    _add_table(
        doc,
        ["参数", "最终值", "解释"],
        [
            ["gap", "1", "与K、M和强正则化共同选择，不能单独解释"],
            ["K", "5", "少量高质量参考，减少噪声参考"],
            ["β", "0.041", "形状权重约0.959，绝对水平只作辅助"],
            ["p", "0.654", "控制距离加权强度"],
            ["M", "1", "每个差值区间只保留少量代表性配对"],
            ["Ridge alpha", "100", "小样本下使用较强正则化"],
        ],
        [1900, 1500, 5960],
    )
    _add_figure(
        doc,
        FIG_SSA,
        "图1  三个随机种子的SSA验证集收敛过程",
        oral_prompt="说明这些参数由验证集搜索得到，不是根据测试结果手工挑选。",
        width=6.0,
    )

    doc.add_page_break()
    _add_heading(doc, "七、7月30日：统一正式环比数据并重新运行", 1)
    _add_body(
        doc,
        "在优化过程中，我与杜艾佳进一步确认了数据口径。仓库早期同时存在同比、"
        "环比、旧版预制特征和不同时间划分；旧传统模型只有30个测试样本，"
        "不能直接与新实验比较。"
    )
    _add_body(
        doc,
        "项目组最终统一为：只使用actual环比序列，连续过去12个月预测下一个月，"
        "训练、验证、测试按50、45、47的固定时间区间划分。"
        "在统一状态通过检查后，我重新运行了正式环比孪生实验。"
    )
    _add_heading(doc, "7.1 核心结果", 2)
    _add_table(
        doc,
        ["模型", "验证MAE", "验证RMSE", "测试MAE", "测试RMSE"],
        [
            ["单光储备池", "0.386178", "0.487698", "0.311511", "0.397189"],
            ["未优化孪生", "0.370512", "0.467292", "0.319792", "0.406210"],
            ["SSA优化孪生", "0.353049", "0.446791", "0.303929", "0.400142"],
        ],
        [2500, 1715, 1715, 1715, 1715],
    )
    _add_body(
        doc,
        "这些指标来自47个测试月份的逐月预测值与对应actual真实值。"
        "MAE是绝对误差平均值，反映一般月份的平均误差；"
        "RMSE对大误差平方后再平均，因此对少数异常月份更敏感。"
    )
    _add_callout(
        doc,
        "核心对比：",
        "SSA孪生测试MAE从0.311511降到0.303929，下降2.43%；"
        "测试RMSE从0.397189升到0.400142，上升0.74%。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )
    _add_body(
        doc,
        "MAE改善说明孪生参考机制在多数月份起到作用；RMSE略差说明仍存在少数较大误差。"
        "因此当前主方案保留单个SSA优化孪生模型，但结论只能是局部改善，不能宣称全面胜出。"
    )
    _add_figure(
        doc,
        FIG_KEY_PRED,
        "图2  47个测试月份的真实值、单光预测与孪生预测曲线",
        oral_prompt="重点指出三条曲线的转折跟随情况，以及少数月份的大误差导致RMSE没有改善。",
        width=6.25,
    )

    doc.add_page_break()
    _add_heading(doc, "八、7月31日：继续优化及为什么放弃", 1)
    _add_heading(doc, "8.1 配对结构优化", 2)
    _add_body(
        doc,
        "我继续从孪生模型最核心的样本配对入手，尝试使用全部合法训练对、"
        "目标与差值区间平衡加权、镜像样本、零截距反对称回归，"
        "以及均值、中位数和截尾均值等聚合方法。"
    )
    _add_body(
        doc,
        "选择这些方法，是因为不同目标产生的配对数量不均衡，模型可能被少数目标或"
        "差值范围主导；镜像和反对称约束则希望让“i相对j”和“j相对i”保持一致。"
    )
    _add_table(
        doc,
        ["模型", "验证MAE", "验证RMSE", "测试MAE", "测试RMSE"],
        [
            ["单光储备池", "0.386178", "0.487698", "0.311511", "0.397189"],
            ["当前SSA孪生", "0.353049", "0.446791", "0.303929", "0.400142"],
            ["配对结构优化", "0.309749", "0.406930", "0.319753", "0.422209"],
        ],
        [2500, 1715, 1715, 1715, 1715],
    )
    _add_body(
        doc,
        "配对结构优化把验证RMSE降低到0.406930，但测试RMSE上升到0.422209，"
        "比单光高6.30%。因此该方案作为负向消融保留，没有替换主模型。"
    )
    _add_body(doc, "我推测失败原因包括：")
    for item in (
        "训练对数量增加，但仍来自相同50个窗口，样本之间高度相关；",
        "镜像和严格反对称在数学上合理，但真实CPI变化不一定完全对称；",
        "模型适应了2018—2022验证期，却没有适应2022—2026测试期；",
        "训练参考池只来自2014—2018，与测试期经济状态存在时间距离。",
    ):
        _add_bullet(doc, item)
    _add_figure(
        doc,
        FIG_PAIR,
        "图3  配对结构候选方案的验证集比较",
        oral_prompt="说明验证集明显改善并不等于未来测试期一定改善，这是本阶段最重要的负结果。",
        width=6.15,
    )

    _add_heading(doc, "8.2 SSA集成和残差式修正", 2)
    _add_body(
        doc,
        "考虑到SSA受随机种子影响，我将三个种子的最佳模型等权平均，"
        "希望通过集成降低单个配置的方差。验证RMSE由0.446791小幅降到0.444837，"
        "但测试RMSE变为0.404185，仍未超过单光。"
    )
    _add_body(
        doc,
        "我认为集成失败的主要原因是三个成员使用相同训练窗口、相同50维状态和相似结构，"
        "错误高度相关。对高度相关模型取平均，降低方差的作用有限。"
    )
    _add_body(
        doc,
        "随后我尝试“单光预测 + λ×(孪生预测−单光预测)”的残差修正，"
        "希望单光提供稳定基础、孪生只负责修正。验证集选择λ=1，"
        "说明验证期完全偏向孪生集成，单光没有提供额外修正价值。"
        "残差方案测试RMSE仍为0.404185，因此也没有采用。"
    )
    _add_figure(
        doc,
        FIG_ENSEMBLE,
        "图4  单光、最佳单个SSA、SSA集成与残差修正的测试指标",
        oral_prompt="说明后续方法虽然有合理动机，但没有形成稳定的测试RMSE提升，所以没有替换主方案。",
        width=6.15,
    )

    doc.add_page_break()
    _add_heading(doc, "九、失败原因的综合判断", 1)
    _add_body(
        doc,
        "综合gap、100维特征、配对结构和集成实验，我认为优化失败不是单一参数造成的，"
        "而是以下因素共同作用："
    )
    failure_rows = [
        ["原始样本过少", "只有50个训练窗口，难以支持高维和复杂结构"],
        ["配对并非独立样本", "组合可增加行数，但无法增加新的经济时期"],
        ["窗口高度重叠", "相邻12个月窗口有大量共同月份，有效信息增量有限"],
        ["时间分布变化", "2014—2018训练参考与2022—2026测试阶段的经济环境不同"],
        ["验证区间单一", "优化算法可能适配45个月验证期，而非一般未来时期"],
        ["模型误差相关", "SSA成员共享数据、状态和结构，集成难以显著降低方差"],
        ["环比信号较弱", "数值集中在100附近，短期波动和噪声占比较高"],
        ["状态为冻结映射", "光储备池状态并非专门针对每个孪生目标重新设计"],
    ]
    _add_table(doc, ["因素", "可能影响"], failure_rows, [2500, 6860])
    _add_callout(
        doc,
        "最重要的判断：",
        "当前主要瓶颈已经从“代码和流程是否跑通”转为“模型能否跨时间阶段泛化”。",
        fill=PALE_GOLD,
        accent="7A5A00",
    )

    _add_heading(doc, "十、传统模型补充对照", 1)
    _add_body(
        doc,
        "为了判断结果所处水平，我还按照完全相同的50/45/47时间划分，"
        "统一运行Ridge、SVR、随机森林、梯度提升等传统模型。"
        "传统模型只作为补充参照，核心任务仍是共享光储备池状态下的单光与孪生比较。"
    )
    _add_table(
        doc,
        ["模型", "测试MAE", "测试RMSE"],
        [
            ["Gradient Boosting", "0.305704", "0.382399"],
            ["Ridge", "0.296734", "0.395175"],
            ["单光储备池", "0.311511", "0.397189"],
            ["SSA孪生光储备池", "0.303929", "0.400142"],
            ["随机森林", "0.327827", "0.402133"],
            ["未优化孪生", "0.319792", "0.406210"],
        ],
        [4360, 2500, 2500],
    )
    _add_figure(
        doc,
        FIG_ALL_MODELS,
        "图5  相同数据口径下全部模型的测试MAE和RMSE",
        oral_prompt="若老师询问传统模型再展示。说明小样本下，强正则化传统模型仍具有竞争力。",
        width=6.15,
    )

    doc.add_page_break()
    _add_heading(doc, "十一、除模型优化外完成的工作", 1)
    other_work = (
        "核对数据口径，明确最终任务是直接预测环比CPI，不再混用同比结果；",
        "统一训练、验证、测试月份，使单光、孪生和传统模型能够直接比较；",
        "审计MATLAB生成的光储备池状态、目标值、样本编号和共享掩码；",
        "建立严格封闭50参考池及防数据泄漏检查；",
        "补充MAT状态文件读取兼容，避免只有npy文件时才能运行；",
        "输出逐月预测、MAE/RMSE表、SSA收敛曲线、残差图和实验清单；",
        "完成传统模型统一对照和结果汇总；",
        "补充自动测试，目前仓库完整测试35项全部通过；",
        "记录每个失败方案，避免后续重复尝试同一无效方向。",
    )
    for item in other_work:
        _add_bullet(doc, item)

    _add_heading(doc, "十二、最终结论与建议", 1)
    _add_body(
        doc,
        "从上周日到现在，我完成的不是单次参数搜索，而是一个从公平性定义、"
        "结构消融、自动优化到失败分析的完整过程。"
    )
    _add_body(doc, "最终保留的方案是：")
    for item in (
        "严格封闭50个原始训练窗口；",
        "与单光共享同一套50维光储备池状态；",
        "使用50维状态差，不采用100维拼接；",
        "参考距离以变化形状为主，绝对水平只占很小权重；",
        "使用SSA选择gap、K、β、p和M；",
        "使用较强Ridge正则化控制小样本过拟合。",
    ):
        _add_bullet(doc, item)
    _add_callout(
        doc,
        "最终结论：",
        "孪生光储备池已经完整跑通，并使测试MAE相对单光下降2.43%；"
        "但测试RMSE上升0.74%，说明孪生历史参考机制具有一定价值，"
        "但尚未形成全面、稳定的优势。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )
    _add_body(
        doc,
        "目前最主要的问题是50个训练窗口过少、组合样本高度相关，"
        "以及训练参考期和测试期存在明显时间分布变化。"
    )
    _add_body(
        doc,
        "如果项目继续，我建议停止根据同一测试区间继续调参，冻结当前主方案，"
        "采用多个滚动时间区间验证，并使用2026年5月之后的新月份作为真正未见数据。"
    )

    _add_heading(doc, "十三、结束语（可直接照读）", 1)
    _add_callout(
        doc,
        "",
        "以上是我从上周日到现在完成的主要工作。我的阶段成果不是证明孪生模型"
        "已经全面超过单光模型，而是完成了孪生光储备池联动、建立了公平的小样本实验，"
        "系统尝试并排除了多种优化方向，同时明确了当前模型的有效部分和主要瓶颈。"
        "目前SSA孪生模型在MAE上取得了改善，下一步重点应从继续扩大参数搜索，"
        "转向提高跨时间阶段的泛化能力。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    doc.add_page_break()
    _add_heading(doc, "附录：老师可能追问的问题", 1)
    qa = [
        (
            "为什么孪生训练对超过50还算严格小样本？",
            "因为所有训练对都由固定50个原始窗口组合，没有访问额外月份。配对行数增加不等于原始数据量增加。",
        ),
        (
            "为什么最终看MAE和RMSE会得到不同结论？",
            "MAE衡量一般月份平均误差，RMSE放大少数大误差。孪生改善多数月份，但大误差控制不足。",
        ),
        (
            "为什么正式主方案保留单个SSA，而没有使用后来验证更好的配对模型？",
            "配对模型验证RMSE更低，但测试明显退化，缺少跨时间稳定性；且测试区间已经被查看，只能作为探索性负结果。",
        ),
        (
            "没有MATLAB是否影响你的工作？",
            "状态由负责单光的同学统一生成。我负责完整审计并在同一状态上训练孪生模型，反而保证了两者共享状态的公平性。",
        ),
        (
            "孪生模型最终是否更好？",
            "在测试MAE上更好2.43%，在RMSE上差0.74%，因此是局部改善，不能表述为全面优于单光。",
        ),
        (
            "下一步最值得做什么？",
            "冻结当前方案，使用滚动时间验证和真正未来的新月份评价，而不是继续根据同一个测试集调参。",
        ),
    ]
    for question, answer in qa:
        _add_body(doc, f"问：{question}", bold_prefix="问：", keep_with_next=True)
        _add_body(doc, f"答：{answer}", bold_prefix="答：")

    _configure_sections(doc)
    for section in doc.sections:
        _add_running_furniture(section)

    doc.core_properties.title = "孪生光储备池小样本CPI预测最终口头汇报报告"
    doc.core_properties.subject = "陈中远项目阶段总结"
    doc.core_properties.author = "陈中远"
    doc.core_properties.keywords = "CPI, 光储备池, 孪生网络, 小样本, SSA"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
